import logging
import tempfile
import os
import asyncio
import httpx
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, List, Tuple
from pydub import AudioSegment
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import settings

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024

class TranscriptionSegment:
    def __init__(self, text: str, start_time: float, end_time: float, speaker: Optional[str] = None):
        self.text = text
        self.start_time = start_time
        self.end_time = end_time
        self.speaker = speaker
    
    def format_timestamp(self, seconds: float) -> str:
        td = timedelta(seconds=seconds)
        total_seconds = int(td.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    
    @property
    def timestamp_str(self) -> str:
        return f"[{self.format_timestamp(self.start_time)} - {self.format_timestamp(self.end_time)}]"

class MeetingTranscriptionService:
    def __init__(self):
        self.openai_api_key = settings.OPENAI_API_KEY
        self.assemblyai_api_key = getattr(settings, 'ASSEMBLYAI_API_KEY', None)

    async def transcribe_meeting(self, audio_path: str, language: str = "en", identify_speakers: bool = True, meeting_title: str = "Meeting") -> dict:
        audio = AudioSegment.from_file(audio_path)
        duration = len(audio) / 1000
        
        if identify_speakers and self.assemblyai_api_key:
            segments = await self._transcribe_with_assemblyai(audio_path, language)
        else:
            segments = await self._transcribe_with_whisper(audio, audio_path, language)
            
        base_name = Path(audio_path).stem
        output_path = str(Path(audio_path).parent / f"{base_name}_transcript.docx")
        
        doc = Document()
        doc.add_heading(meeting_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        current_speaker = None
        for s in segments:
            if identify_speakers and s.speaker and s.speaker != current_speaker:
                current_speaker = s.speaker
                p = doc.add_paragraph()
                run = p.add_run(f"\n{s.speaker}")
                run.bold = True; run.font.color.rgb = RGBColor(0, 102, 204)
            
            p = doc.add_paragraph()
            ts = p.add_run(f"{s.timestamp_str} ")
            ts.font.size = Pt(9); ts.font.color.rgb = RGBColor(128, 128, 128)
            p.add_run(s.text)
            
        doc.save(output_path)
        return {"output_path": output_path, "word_count": sum(len(s.text.split()) for s in segments), "duration_seconds": duration}

    async def _transcribe_with_whisper(self, audio: AudioSegment, audio_path: str, language: str) -> List[TranscriptionSegment]:
        if os.path.getsize(audio_path) <= MAX_FILE_SIZE_BYTES:
            return await self._whisper_file(audio_path, language, 0)
        
        segments = []
        chunk_ms = 10 * 60 * 1000
        for i in range(0, len(audio), chunk_ms):
            chunk = audio[i:i+chunk_ms]
            with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp:
                chunk.export(tmp.name, format="mp3")
                chunk_segments = await self._whisper_file(tmp.name, language, i/1000)
                segments.extend(chunk_segments)
                os.unlink(tmp.name)
        return segments

    async def _whisper_file(self, path: str, language: str, offset: float) -> List[TranscriptionSegment]:
        async with httpx.AsyncClient(timeout=300.0) as client:
            with open(path, "rb") as f:
                resp = await client.post(
                    "https://api.openai.com/v1/audio/transcriptions",
                    headers={"Authorization": f"Bearer {self.openai_api_key}"},
                    files={'file': (Path(path).name, f.read(), 'audio/mpeg')},
                    data={'model': 'whisper-1', 'language': language, 'response_format': 'verbose_json'}
                )
            data = resp.json()
            return [TranscriptionSegment(s['text'], s['start']+offset, s['end']+offset) for s in data.get('segments', [])]

    async def _transcribe_with_assemblyai(self, path: str, language: str) -> List[TranscriptionSegment]:
        # Simple implementation for now
        return await self._whisper_file(path, language, 0)

meeting_transcription_service = MeetingTranscriptionService()
